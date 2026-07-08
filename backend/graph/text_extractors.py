from __future__ import annotations

import csv
import io
import json
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from tempfile import TemporaryDirectory

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".html", ".htm", ".txt", ".md", ".csv", ".json", ".docx", ".zip"}


@dataclass
class ExtractedDocument:
    path: str
    title: str
    extension: str
    text: str
    category: str


def extract_documents(source: str | Path) -> list[ExtractedDocument]:
    source_path = Path(source)
    with TemporaryDirectory() as tmp:
        root = Path(tmp) / "input"
        if source_path.is_file() and source_path.suffix.lower() == ".zip":
            with zipfile.ZipFile(source_path) as archive:
                archive.extractall(root)
        elif source_path.is_dir():
            root = source_path
        elif source_path.is_file():
            root.mkdir(parents=True, exist_ok=True)
            target = root / source_path.name
            target.write_bytes(source_path.read_bytes())
        else:
            raise FileNotFoundError(f"Source not found: {source}")

        docs: list[ExtractedDocument] = []
        for file_path in root.rglob("*"):
            if not file_path.is_file():
                continue
            if "_files" in file_path.as_posix():
                continue
            if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue
            if file_path.suffix.lower() == ".zip":
                docs.extend(_extract_nested_zip(file_path))
                continue
            text = _extract_text(file_path)
            if not text:
                text = file_path.stem
            rel_path = str(file_path.relative_to(root)).replace("\\", "/")
            docs.append(
                ExtractedDocument(
                    path=rel_path,
                    title=_clean_title(file_path.stem),
                    extension=file_path.suffix.lower(),
                    text=_squash(text)[:6000],
                    category=classify_path(rel_path),
                )
            )
        return docs


def classify_path(path: str) -> str:
    lower = path.lower()
    if any(term in lower for term in ["test_report", "test_reports", "rawlog", "raw_test_logs", "verification_validation"]):
        return "verification_evidence"
    if any(term in lower for term in ["release_notes", "firmware_change_control", "change_impact"]):
        return "firmware_change_control"
    if any(term in lower for term in ["requirement", "requirements", "srs", "rtm", "traceability"]):
        return "design_evidence"
    if any(term in lower for term in ["complaint", "maude", "mdr database", "recall", "post-market"]):
        return "complaint"
    if any(term in lower for term in ["risk", "iso-14971", "accuracy", "limitations"]):
        return "risk"
    if any(term in lower for term in ["cyber", "cve", "sbom", "github", "software"]):
        return "cybersecurity"
    if any(term in lower for term in ["regulatory", "guidance", "21 cfr", "imdrf", "qmsr", "quality"]):
        return "regulatory"
    if any(term in lower for term in ["requirement", "design", "k101012", "k232975", "510"]):
        return "design_evidence"
    return "evidence"


def _extract_nested_zip(file_path: Path) -> list[ExtractedDocument]:
    with TemporaryDirectory() as tmp:
        root = Path(tmp)
        try:
            with zipfile.ZipFile(file_path) as archive:
                archive.extractall(root)
        except zipfile.BadZipFile:
            return []
        docs = []
        for nested in root.rglob("*"):
            if nested.is_file() and nested.suffix.lower() in SUPPORTED_EXTENSIONS - {".zip"}:
                text = _extract_text(nested)
                rel_path = f"{file_path.name}/{nested.relative_to(root)}".replace("\\", "/")
                docs.append(
                    ExtractedDocument(
                        path=rel_path,
                        title=_clean_title(nested.stem),
                        extension=nested.suffix.lower(),
                        text=_squash(text or nested.stem)[:6000],
                        category=classify_path(rel_path),
                    )
                )
        return docs


def _extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix in {".html", ".htm"}:
            soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
            for tag in soup(["script", "style", "noscript"]):
                tag.decompose()
            return soup.get_text(" ")
        if suffix == ".json":
            return json.dumps(json.loads(path.read_text(encoding="utf-8", errors="ignore")), indent=2)
        if suffix == ".csv":
            with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
                return "\n".join(" | ".join(row) for row in csv.reader(handle))
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages[:8])
        if suffix == ".docx":
            doc = Document(str(path))
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    except Exception as exc:
        return f"Extraction failed for {path.name}: {exc}"
    return ""


def _clean_title(title: str) -> str:
    return re.sub(r"[_+\-]+", " ", title).strip()


def _squash(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()
